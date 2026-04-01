#!/usr/bin/env python3
"""
Extract phonics data and clipart images from clipart assets.docx.
Run from the worksheetgenerator/ directory:
    python3 scripts/extract_assets.py
"""

import sys
import os
import zipfile
import json
import re
from pathlib import Path

try:
    from docx import Document
    from docx.oxml.ns import qn
except ImportError:
    print("ERROR: python-docx not found.")
    print("Create a venv and install it: python3 -m venv /tmp/env && /tmp/env/bin/pip install python-docx")
    sys.exit(1)

SCRIPT_DIR = Path(__file__).parent
ROOT = SCRIPT_DIR.parent
SCHOOL = ROOT.parent  # .../claude/school
ASSETS_DOC = SCHOOL / "assets" / "clipart assets.docx"
IMAGES_OUT = ROOT / "public" / "images"
DATA_OUT = ROOT / "src" / "data" / "phonicsData.js"

# ── Actual colours from doc inspection ──────────────────────────────────────
SECTION_COLORS = {
    'FF6B6B': 'CVC',
    '7C3AED': 'CVCe',
    '38B2AC': 'Beginning Blends',
    '10B981': 'Ending Blends',
    'D97706': 'Vowel Teams & Diphthongs',
}

TEMPLATE_TYPE = {
    'CVC':                       'cvc',
    'CVCe':                      'blends',
    'Beginning Blends':          'blends',
    'Ending Blends':             'blends',
    'Vowel Teams & Diphthongs':  'blends',
}

# Known top-level section header phrases (lower-cased)
SECTION_HEADERS = [
    'cvc word families',
    'cvce word families',
    'beginning blends',
    'ending blends',
    'vowel teams',
    'diphthong',
    'r-controlled',
]

# Known family / sub-section header patterns
FAMILY_HEADER_RE = re.compile(
    r'^(short [aeiou]|long [aeiou]|l blends|r blends|s blends|3-letter|nasal|stop|mixed|'
    r'r-controlled|diphthong|long vowel)', re.IGNORECASE
)

def slugify(text):
    return re.sub(r'[^a-z0-9]+', '-', text.lower()).strip('-')

def get_first_color(para):
    """Return the hex colour string of the first coloured run, or None."""
    for r in para.runs:
        try:
            if r.font.color.type:
                return str(r.font.color.rgb)
        except Exception:
            pass
    return None

def build_word_image_map(doc):
    """Return {word: image_filename} for every clip_ bookmark in the doc."""
    body_children = list(doc.element.body)

    # rId → filename inside word/media/
    img_rels = {}
    for rId, rel in doc.part.rels.items():
        if 'image' in rel.reltype:
            img_rels[rId] = Path(rel.target_ref).name

    def get_image_rId(el):
        blip = el.find('.//' + qn('a:blip'))
        return blip.get(qn('r:embed')) if blip is not None else None

    word_image = {}
    for i, child in enumerate(body_children[:-1]):
        rId = get_image_rId(child)
        if rId:
            nxt = body_children[i + 1]
            for b in nxt.findall('.//' + qn('w:bookmarkStart')):
                name = b.get(qn('w:name'), '')
                if name.startswith('clip_'):
                    word = name[5:]
                    word_image[word] = img_rels.get(rId, '')
    return word_image

def collect_clipart_words(doc):
    """Return set of words wrapped in clip_ hyperlinks."""
    result = set()
    for p in doc.paragraphs:
        for hl in p._element.findall('.//' + qn('w:hyperlink')):
            anchor = hl.get(qn('w:anchor'), '')
            if anchor.startswith('clip_'):
                word = ''.join(t.text or '' for t in hl.iter(qn('w:t'))).strip()
                result.add(word)
    return result

def parse_word_line(text, starts_with_dash):
    """
    Parse a phonics word line and return (rime_token, [words]).

    Two formats exist:
      1. CVC/CVCe/Ending Blends:  '-at  bat  cat  fat...'
      2. Beginning Blends/Vowel Teams:  'bl-  black  blade...' / 'ai  bait...'
    """
    parts = text.split()
    if len(parts) < 2:
        return None, []

    rime = parts[0]
    words = [w.strip().rstrip('.,') for w in parts[1:] if w.strip()]
    return rime, words

def is_word_list_line(text, color):
    """
    Return True if this non-empty paragraph looks like a word-list line.

    Three formats:
      1. CVC / Ending Blends: '-at  bat  cat  fat…'   (first token starts with '-')
      2. Beginning Blends:    'bl-  black  blade…'    (first token ends with '-')
      3. Vowel Teams:         'ai  bait  brain…'      (first token is 1-4 lowercase
                                                       letters; ALL remaining tokens
                                                       must be all-lowercase to rule
                                                       out section headers like
                                                       'CVC Word Families')
    """
    if not color or color not in SECTION_COLORS:
        return False
    parts = text.split()
    if len(parts) < 2:
        return False
    first = parts[0]
    rest = parts[1:]

    # Format 1: '-rime  word1  word2…'
    if first.startswith('-') and len(first) >= 2:
        return True
    # Format 2: blend prefix 'bl-', 'str-', etc.
    if re.match(r'^[a-zA-Z]{1,4}-$', first):
        return True
    # Format 3: vowel team / r-controlled pattern.
    # The first token is 1-4 lowercase letters AND every subsequent token is
    # all-lowercase (real words, not title-case header words like "Word Families").
    if (re.match(r'^[a-z]{1,4}$', first)
            and len(rest) >= 2
            and all(w == w.lower() for w in rest)):
        return True
    return False

def parse_structure(doc, clipart_words):
    """
    Walk paragraphs (stopping at 'Clipart Library') and return list of pattern dicts.
    """
    current_category = None
    current_family = None
    patterns = []

    for para in doc.paragraphs:
        text = para.text.strip()
        if not text:
            continue

        # Stop at the clipart gallery section
        if 'Clipart Library' in text:
            break

        color = get_first_color(para)
        tl = text.lower()

        # ── Classify this paragraph ──────────────────────────────────────
        if is_word_list_line(text, color):
            # It's a word list line
            rime, word_list = parse_word_line(text, text.startswith('-'))
            words = []
            for w in word_list:
                w_clean = w.strip().rstrip('.,')
                if w_clean:
                    words.append({'word': w_clean, 'hasClipart': w_clean in clipart_words})

            if any(e['hasClipart'] for e in words):
                category = current_category or 'Other'
                pat_id = slugify(f"{category} {current_family or ''} {rime}")
                family_str = current_family or ''
                display = f"{rime}"
                if family_str:
                    display = f"{family_str}  {rime}"

                patterns.append({
                    'id': pat_id,
                    'displayName': display,
                    'category': category,
                    'family': family_str,
                    'templateType': TEMPLATE_TYPE.get(category, 'blends'),
                    'rime': rime,
                    'words': words,
                })

        elif color and color in SECTION_COLORS:
            # It's a coloured header — figure out if top-level section or family
            if any(kw in tl for kw in SECTION_HEADERS):
                current_category = SECTION_COLORS[color]
                current_family = None
            elif FAMILY_HEADER_RE.match(tl):
                current_family = text
            else:
                # Could be a sub-family like "Short A Families" or "Long Vowel Teams"
                current_family = text

    return patterns

def extract_images(docx_path, word_image_map, out_dir):
    """Copy word clipart images out of the docx."""
    with zipfile.ZipFile(docx_path) as z:
        media_files = {Path(f).name: f for f in z.namelist() if f.startswith('word/media/')}
        copied = 0
        for word, img_file in word_image_map.items():
            if img_file and img_file in media_files:
                data = z.read(media_files[img_file])
                ext = Path(img_file).suffix
                out_path = out_dir / f"word_{word}{ext}"
                out_path.write_bytes(data)
                copied += 1
    return copied

def main():
    print(f"Reading: {ASSETS_DOC}")
    if not ASSETS_DOC.exists():
        print(f"ERROR: File not found: {ASSETS_DOC}")
        sys.exit(1)

    doc = Document(str(ASSETS_DOC))

    print("Building word→image map…")
    word_image_map = build_word_image_map(doc)
    print(f"  {len(word_image_map)} words with clipart images")

    print("Collecting clipart words…")
    clipart_words = collect_clipart_words(doc)
    print(f"  {len(clipart_words)} clipart words found")

    print("Parsing phonics structure…")
    patterns = parse_structure(doc, clipart_words)
    print(f"  {len(patterns)} patterns with clipart words found")

    # Attach image filenames to word entries
    for pat in patterns:
        for w in pat['words']:
            if w['hasClipart']:
                orig_img = word_image_map.get(w['word'], '')
                if orig_img:
                    ext = Path(orig_img).suffix
                    w['image'] = f"word_{w['word']}{ext}"

    # Summary by category
    from collections import Counter
    cats = Counter(p['category'] for p in patterns)
    for cat, n in sorted(cats.items()):
        print(f"    {cat}: {n} patterns")

    print(f"Extracting images to {IMAGES_OUT}…")
    IMAGES_OUT.mkdir(parents=True, exist_ok=True)
    n = extract_images(str(ASSETS_DOC), word_image_map, IMAGES_OUT)
    print(f"  {n} images extracted")

    # Write JS data module
    DATA_OUT.parent.mkdir(parents=True, exist_ok=True)
    js = "// AUTO-GENERATED — run scripts/extract_assets.py to regenerate\n"
    js += "export const phonicsData = " + json.dumps({'patterns': patterns}, indent=2) + ";\n"
    DATA_OUT.write_text(js)
    print(f"Data written to {DATA_OUT}")
    print("Done.")

if __name__ == '__main__':
    main()
